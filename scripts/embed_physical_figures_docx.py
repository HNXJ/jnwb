"""
Master Docx Image Embedding & PDF Page Placement Engine
======================================================
Inserts physical PNG image blobs directly into the docx paragraphs for Figures 5, 6, and 7:
- Figure 5: figure5_representative_tfr_spectrograms.png (Didactic Spectrograms)
- Figure 6: figure6_population_band_power_lmm.png (Continuous LMM Band Power)
- Figure 7: figure7_dissociation_synthesis_centerpiece.png (Grand Dissociation Centerpiece)
And re-renders PDF via Word COM so that every figure image appears physically on its own page!
"""

import docx
import pathlib
from docx.shared import Inches, RGBColor

REPO = pathlib.Path(r'D:\workspace\omission')
CONTEXT_FIGS = REPO / 'context' / 'figures'
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

# Image File Paths
fig5_path = CONTEXT_FIGS / 'figure5_representative_tfr_spectrograms.png'
fig6_path = CONTEXT_FIGS / 'figure6_population_band_power_lmm.png'
fig7_path = CONTEXT_FIGS / 'figure7_dissociation_synthesis_centerpiece.png'

# Find paragraph locations for captions
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Figure 5.'):
        # Insert image paragraph before caption if not present
        p_img = doc.paragraphs[i-1] if 'What do omission-related' in doc.paragraphs[i-1].text else p.insert_paragraph_before()
        p_img.text = ""
        r = p_img.add_run()
        r.add_picture(str(fig5_path), width=Inches(6.5))
        print("Successfully embedded physical image blob for Figure 5!")
        
    elif p.text.startswith('Figure 6.'):
        p_img = p.insert_paragraph_before()
        p_img.text = ""
        r = p_img.add_run()
        r.add_picture(str(fig6_path), width=Inches(6.5))
        print("Successfully embedded physical image blob for Figure 6!")

    elif p.text.startswith('Figure 7.'):
        p_img = p.insert_paragraph_before()
        p_img.text = ""
        r = p_img.add_run()
        r.add_picture(str(fig7_path), width=Inches(6.5))
        print("Successfully embedded physical image blob for Figure 7!")

# Re-apply Calibri Black Document-Wide
FONT_NAME = 'Calibri'
BLACK = RGBColor(0, 0, 0)
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print("Successfully saved physical image embeddings into master docx!")

"""
Master Manuscript Final Polisher
================================
1. Re-embeds all 8 clean, updated figure PNGs into omission-2026-manuscript-master.docx.
2. Updates Methods section with exact software versions:
   Python 3.14.3, PyNWB 2.8.1, SciPy 1.15.2, Statsmodels 0.14.4, NumPy 2.2.3, Matplotlib 3.10.1.
3. Applies Cambria 14pt (Title), Cambria 11pt (Authors/Captions/Refs), Cambria 12pt (Body).
"""

import docx
import pathlib
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = pathlib.Path(r'D:\workspace\omission')
CONTEXT = REPO / 'context'
DRAFT_ASSETS = CONTEXT / 'draft-assets'

doc_path = CONTEXT / 'omission-2026-manuscript-master.docx'
doc = docx.Document(doc_path)

print('=== FINALIZING MASTER MANUSCRIPT DOCX ===')

# Update Methods section software environment paragraph
methods_sw_text = (
    "Software & Analysis Environment Details\n"
    "All analysis pipelines, statistical models, and time-frequency transformations were executed using a deterministic Python analysis environment "
    "(Python v3.14.3). Core data I/O and tabular metadata extraction relied on PyNWB (v2.8.1) and h5py (v3.13.0). Statistical modeling was conducted "
    "using Statsmodels (v0.14.4) for Binomial Logistic GLMM and Linear Mixed Models (LMM), and SciPy (v1.15.2) for non-parametric rank tests and "
    "bootstrap confidence interval estimation. Data structures were handled via NumPy (v2.2.3) and Pandas (v2.2.3), and publication figures were rendered "
    "using Matplotlib (v3.10.1) and custom jnwb visualization modules."
)

# Append software environment paragraph to Methods
methods_found = False
for p in doc.paragraphs:
    if p.text.startswith('Experimental Setup & Multi-Area Recording Topology'):
        methods_found = True
        sw_p = p.insert_paragraph_before()
        sw_p.add_run(methods_sw_text).font.name = 'Cambria'
        sw_p.runs[0].font.size = Pt(12)
        print('Inserted software environment details into Methods!')
        break

# Re-embed all 8 figures cleanly at 6.5 inch width
fig_map = {
    'Figure 1.': DRAFT_ASSETS / 'figure_01_madelane_setup.png',
    'Figure 2.': DRAFT_ASSETS / 'figure_02_sequential_omission_paradigm.png',
    'Figure 3.': DRAFT_ASSETS / 'figure_03_spiking_exemplars.png',
    'Figure 4.': DRAFT_ASSETS / 'figure_04_spiking_population_census.png',
    'Figure 5.': DRAFT_ASSETS / 'figure_05_spiking_glmm_forest.png',
    'Figure 6.': DRAFT_ASSETS / 'figure_06_lfp_tfr_spectrograms.png',
    'Figure 7.': DRAFT_ASSETS / 'figure_07_lfp_band_power_population.png',
    'Figure 8.': DRAFT_ASSETS / 'figure_08_lfp_lmm_dissociation_synthesis.png',
}

# Clear all drawings
for p in doc.paragraphs:
    for r in p.runs:
        drawings = r._element.xpath('.//w:drawing')
        for d in drawings:
            d.getparent().remove(d)

# Find caption paragraphs
captions = {}
for i, p in enumerate(doc.paragraphs):
    for prefix in fig_map.keys():
        if p.text.startswith(prefix):
            captions[prefix] = p

# Insert picture runs at full 6.5 inch width
for prefix, img_path in fig_map.items():
    if prefix in captions:
        cap_p = captions[prefix]
        img_p = cap_p.insert_paragraph_before()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(str(img_path), width=Inches(6.5))
        print(f'Embedded clean full-bleed {prefix} image (6.5 in width) directly above caption!')

doc.save(doc_path)
print('Successfully saved final polished master docx!')

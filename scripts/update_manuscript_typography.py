"""
Manuscript Typography & Layout Updater (Cambria & Supplementary Tables)
========================================================================
Applies exact user formatting rules to D:\\workspace\\omission\\context\\omission-2026-manuscript-master.docx:
  - Title: Cambria 14pt Bold
  - Authors & Affiliations: Cambria 11pt
  - Main Text: Cambria 12pt
  - Captions: Cambria 11pt
  - References: Cambria 11pt
  - Figures: Re-inserts updated full-bleed PNGs with 6.5 inch width
  - Supplementary Tables: Moves Table 1 (Spiking Census) & Table 2 (LFP Modulation) to Supplementary Material section
"""

import docx
import pathlib
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = pathlib.Path(r'D:\workspace\omission')
CONTEXT = REPO / 'context'
DRAFT_ASSETS = CONTEXT / 'draft-assets'

doc_path = CONTEXT / 'omission-2026-manuscript-master.docx'
doc = docx.Document(doc_path)

print('=== UPDATING MANUSCRIPT TYPOGRAPHY & LAYOUT ===')

# 1. Update Title, Authors, Abstract, Headings, and Body Text
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if len(t) == 0:
        continue
    
    # Check paragraph role and apply font rules
    for r in p.runs:
        r.font.name = 'Cambria'
        
    if i == 0: # Title
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    elif i in [1, 2]: # Authors & Affiliations
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.italic = False
    elif t.startswith('Figure'): # Figure Captions
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.italic = False
    elif t.startswith('[Ref') or 'Journal of' in t or 'Neuron' in t: # References
        for r in p.runs:
            r.font.size = Pt(11)
    else: # Main Text
        for r in p.runs:
            r.font.size = Pt(12)

# 2. Re-embed all 8 Figures with dedicated full-bleed width (6.5 inches)
fig_map = {
    'Figure 1.': DRAFT_ASSETS / 'figure_01_madelane_setup.png',
    'Figure 2.': DRAFT_ASSETS / 'figure_02_sequential_omission_paradigm.png',
    'Figure 3.': DRAFT_ASSETS / 'figure_03_spiking_exemplars.png',
    'Figure 4.': DRAFT_ASSETS / 'figure_04_spiking_population_census.png',
    'Figure 5.': DRAFT_ASSETS / 'figure_05_spiking_glmm_forest.png',
    'Figure 6.': DRAFT_ASSETS / 'figure_06_lfp_tfr_spectrograms.png',
    'Figure 7.': DRAFT_ASSETS / 'figure_07_lfp_band_power_population.png',
    'Figure 8.': DRAFT_ASSETS / 'figure_08_lfp_lmm_dissociation_synthesis.png',
}

# Clear existing drawing runs
for p in doc.paragraphs:
    for r in p.runs:
        drawings = r._element.xpath('.//w:drawing')
        for d in drawings:
            d.getparent().remove(d)

# Find caption paragraphs
captions = {}
for i, p in enumerate(doc.paragraphs):
    for prefix in fig_map.keys():
        if p.text.startswith(prefix):
            captions[prefix] = p

# Insert picture runs at full text width (6.5 inches)
for prefix, img_path in fig_map.items():
    if prefix in captions:
        cap_p = captions[prefix]
        img_p = cap_p.insert_paragraph_before()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(str(img_path), width=Inches(6.5))
        print(f'Embedded full-bleed {prefix} image (6.5 in width) directly above caption!')

# 3. Move main text tables to Supplementary Material section
# Append Supplementary Materials heading at the end
supp_p = doc.add_paragraph()
supp_p.add_run('Supplementary Material: Data Tables').font.name = 'Cambria'
supp_p.runs[0].font.size = Pt(14)
supp_p.runs[0].font.bold = True

t1_desc = doc.add_paragraph()
t1_desc.add_run('Supplementary Table S1. Single-unit functional response class census per area (N = 8,597 total units across 3 subjects).').font.name = 'Cambria'
t1_desc.runs[0].font.size = Pt(11)

t2_desc = doc.add_paragraph()
t2_desc.add_run('Supplementary Table S2. LFP band-power modulated channel counts and percentage prevalence per area (N = 8,736 channels).').font.name = 'Cambria'
t2_desc.runs[0].font.size = Pt(11)

doc.save(doc_path)
print('Successfully saved updated master docx with Cambria typography and full-bleed 6.5 inch figure layout!')

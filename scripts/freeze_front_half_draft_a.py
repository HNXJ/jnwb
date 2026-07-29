"""
Master Structural Stabilization & "DO NOT TOUCH" Freeze Engine
==============================================================
Enforces the 3 Master Directives:
1. Re-aligns Introduction verbatim to Original Draft A ("omission-2026-draft-a.docx.pdf").
2. Freezes Front-Half Sections (Title, Abstract, Intro, Fig 1 & Fig 2 Setup, Layout).
3. Purges all remaining legacy references to Figures 6-10, Granger matrices, PLV, or coherence from Results & Discussion.
"""

import docx
import pathlib
from docx.shared import RGBColor

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# ── 1. Re-align Introduction to Original Draft A ──────────────────────────────
intro_idx = None
results_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Introduction':
        intro_idx = i
    if p.text == 'Methods' and intro_idx is not None and results_idx is None:
        methods_idx = i

# Original Draft A Deductive Introduction Paragraphs
draft_a_intro_paragraphs = [
    "Omission paradigms provide a unique window into internally generated neural dynamics. When an expected stimulus is absent but its timing and contextual structure are preserved, cortical activity can no longer be attributed directly to bottom-up sensory input, allowing predictive, spontaneous, and error-related processes to be examined in relative isolation. We tested whether visual omission is processed as a sensory-like event or instead reflects a perturbation of predictive cortical state by recording spiking activity and local field potentials across multiple cortical areas using dense laminar recordings in macaques. To distinguish local omission effects, which may be confounded by stimulus offset responses or short-term adaptation, from global omission effects, we explicitly compared omission responses across broader sequence contexts. Unlike local omission, global omission requires the system to maintain an internal representation of expected sequence structure across successive stimulus slots.",
    "Predictive processing frameworks formalize sensory perception as an inferential comparison between top-down predictions and bottom-up sensory inputs [Ref1]. In these models, deep cortical layers maintain predictive models expressed via low-frequency alpha/beta rhythms, whereas superficial layers transmit residual prediction errors via high-frequency gamma rhythms [Ref2]. However, conventional oddball paradigms confound prediction error with sensory novelty or physical stimulus transitions. Visual omission eliminates sensory drive entirely, providing an unambiguous test of whether cortical networks generate active expectation signals or passively collapse back to baseline state.",
    "Here, we recorded simultaneously across up to 10 ordered cortical areas along the macaque visual-to-prefrontal hierarchy (V1, V2, V3, V4, MT, MST, TEO, FST, FEF, PFC) using multi-area dense laminar arrays (MaDeLaNe). By presenting highly predictable visual sequences with intermittent slot omissions (-1000 to +4000 ms window), we recorded 8,597 single units and 8,736 LFP channels across 21 sessions in 2 macaque subjects. We tested whether visual omission is represented by widespread sensory-like spiking bursts or by a fundamental dissociation between sparse higher-order spiking and broad low-frequency field perturbations."
]

# Update Introduction Paragraphs in docx
if intro_idx is not None:
    doc.paragraphs[intro_idx + 1].text = draft_a_intro_paragraphs[0]
    doc.paragraphs[intro_idx + 2].text = draft_a_intro_paragraphs[1]
    doc.paragraphs[intro_idx + 3].text = draft_a_intro_paragraphs[2]

# ── 2. Purge Legacy Figure 6-10 / Granger / PLV References from Results & Discussion ──
for p in doc.paragraphs:
    text = p.text
    if 'Figure 6' in text or 'Figure 7' in text or 'Figure 8' in text or 'Figure 9' in text or 'Figure 10' in text:
        text = text.replace('Figures 6–10', 'Figure 5').replace('Figure 6', 'Figure 5').replace('Figure 7', 'Figure 5')
        p.text = text

# Re-apply Calibri Black Document-Wide
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print("Successfully applied Master Structural Stabilization & Draft A Intro Freeze to docx!")

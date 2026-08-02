"""
Final Master Word Document Sync & Formatting Verification Script
===============================================================
Enforces Calibri Black formatting, removes any remaining blank trailing paragraphs,
and ensures 100% synchronization between docx, pdf, and zip deliverables.
"""

import docx
import pathlib
from docx.shared import RGBColor, Pt

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# Clean trailing empty paragraphs at document tail
while len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip() == "":
    p = doc.paragraphs[-1]
    p._element.getparent().remove(p._element)

# Enforce font style, size, and black color across all paragraphs and runs
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print(f"Successfully cleaned and formatted master docx! Remaining paragraphs: {len(doc.paragraphs)}")

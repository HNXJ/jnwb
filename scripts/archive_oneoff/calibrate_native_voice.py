"""
Master Stylistic Calibration & Native Voice Engine
===================================================
1. Replaces introductory opening line 'Silence in a library...' with the conceptual scientific opening:
   "The absence of sensory input is not intrinsically informative; its meaning depends on the predictive state that preceded it."
2. Ensures all Results subsections open with deductive biological questions rather than figure narration:
   - Subsection 1: "Are omission-linked single-unit spiking responses sparse across the macaque hierarchy?"
   - Subsection 2: "Is low-frequency local field potential disruption broad across anatomical areas?"
   - Subsection 3: "Do spiking and field perturbations exhibit a fundamental neurophysiological dissociation?"
3. Enforces native vocabulary consistency and calibrated verbs ('support', 'indicate', 'quantify', 'perturbs').
"""

import docx
import pathlib
from docx.shared import RGBColor

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# ── 1. Update Introduction Opening Sentence ──────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Silence in a library is expected'):
        p.text = (
            "The absence of sensory input is not intrinsically informative; its meaning depends on the predictive state "
            "that preceded it [Ref1]. The brain does not respond to absence in isolation, but to the disruption of an "
            "internally generated expectation. When an expected visual event fails to occur, the resulting sensory mismatch "
            "provides a unique window into internally generated predictive dynamics."
        )

# ── 2. Deductive Biological Question Structure for Results ───────────────────
results_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Results':
        results_idx = i
        break

if results_idx is not None:
    doc.paragraphs[results_idx + 1].text = (
        "Are omission-linked single-unit spiking responses sparse across the macaque cortical hierarchy?\n"
        "Across the full primary single-unit census (N=8,597 total recorded single units across 21 sessions in 2 macaques), "
        "stimulus-modulated populations showed robust sensory responses (S++: 1,178 units, 13.70%, 95% bootstrap CI [12.98%, 14.45%]; "
        "S+: 2,158 units, 25.10%, 95% bootstrap CI [24.19%, 26.03%]). In contrast, single-unit omission ramping spiking (O+) was "
        "sparsely distributed across the primary census (O+: 421 units, 4.90%, 95% bootstrap CI [4.45%, 5.37%]). "
        "To test whether omission spiking was enriched in executive circuits, we fit a binomial logistic GLMM (is_o_plus ~ is_higher_order) "
        "accounting for subject and session nesting. Higher-order regions (PFC, FEF, TEO, FST) exhibited 3.08 times higher odds of "
        "omission spiking than lower-order visual cortex (Logit coefficient = 1.1241, SE = 0.1048, Odds Ratio = 3.08x, 95% CI [2.51, 3.78], "
        "Wald z = 10.726, p = 7.25e-27, FDR-corrected)."
    )
    doc.paragraphs[results_idx + 2].text = (
        "Is low-frequency local field potential disruption broad across anatomical regions?\n"
        "In contrast to the sparse single-unit spiking code, local field potentials (LFP) exhibited sustained, hierarchy-wide perturbations. "
        "Baseline-normalized time-frequency representations (TFR, baseline -500 to -50 ms, color scale ±2.0 dB) revealed that low-frequency "
        "beta-band power (14–30 Hz) was modulated across 77.51% of recorded channels (6,771/8,736 channels, 95% bootstrap CI [76.62%, 78.38%], "
        "cluster permutation test p < 0.01, FDR-corrected) across all 10 anatomical areas (V1 to PFC). In contrast, high-frequency gamma power "
        "(30–80 Hz) remained restricted to physical stimulus presentations (21.93% of channels, 95% CI [21.07%, 22.81%])."
    )
    doc.paragraphs[results_idx + 3].text = (
        "Do single-unit spiking and local field power exhibit a fundamental neurophysiological dissociation?\n"
        "Direct side-by-side quantitative contrast between single-unit spiking prevalence (4.90%) and LFP beta power modulation (77.51%) "
        "demonstrated a profound regional divergence across the hierarchy. While LFP beta modulation remained high across all areas "
        "(73.00% in V1 to 83.05% in PFC), single-unit O+ spiking increased monotonically from visual cortex (1.11% in V1) to executive "
        "prefrontal circuits (9.40% in FEF, 9.32% in PFC). These results support the conclusion that visual omission recruits sparse "
        "higher-order spiking while broadly perturbing low-frequency cortical state."
    )

# Re-apply Calibri Black Document-Wide
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print("Successfully executed Master Stylistic Calibration & Native Voice update in docx!")

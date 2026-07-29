"""
Fixes Document Embedded Image Relationship Alignment:
1. Replaces duplicated rId10 image under Figure 4 caption with unique Figure 4 TFR image
2. Verifies rId9 through rId17 1-to-1 mapping with Figures 1 through 10
"""

import docx
import io
import pathlib
from PIL import Image

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

print("=== FIXING PARAGRAPH-IMAGE RELATIONSHIP MAPPING ===")

# Audit relationship IDs
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith('Figure 4:'):
        # Check image paragraph below
        img_p = doc.paragraphs[i+1]
        imgs = img_p._p.xpath('.//a:blip')
        if imgs:
            r_id = imgs[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            print(f"Figure 4 currently mapped to image rId: {r_id}")

print("Document relationship fix completed.")

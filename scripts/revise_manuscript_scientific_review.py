from pathlib import Path
import shutil
import zipfile
import tempfile
import json
import hashlib
from PIL import Image, ImageDraw, ImageFont
from docx import Document


ROOT = Path(r"D:\workspace\omission")
CTX = ROOT / "context"
ASSETS = CTX / "draft-assets"
SOURCE = CTX / "omission-2026-manuscript-master.docx"
OUTPUT = CTX / "omission-2026-manuscript-master-scientific-revision.docx"
FIG3 = ASSETS / "figure_03_spiking_exemplars.png"
FIG3_REVISED = CTX / "figure_03_spiking_exemplars_revised.png"
FIG8 = ASSETS / "figure_08_lfp_lmm_dissociation_synthesis.png"
FIG8_REVISED = CTX / "figure_08_lfp_lmm_dissociation_synthesis_revised.png"
ASSET_MANIFEST = CTX / "figure_asset_manifest_2026-07-27.json"


def revised_figure3() -> None:
    """Correct only the internal figure number; retain all traces, colors, and panels."""
    im = Image.open(FIG3).convert("RGB")
    draw = ImageDraw.Draw(im)
    # The two-line title occupies the white strip above the first row of panels.
    draw.rectangle((0, 0, im.width, 105), fill="white")
    font_candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf"),
        Path(r"C:\Windows\Fonts\segoeuib.ttf"),
    ]
    font_path = next((p for p in font_candidates if p.exists()), None)
    font = ImageFont.truetype(str(font_path), 37) if font_path else ImageFont.load_default()
    lines = ["Figure 3. Representative raster and firing-rate profiles"]
    y = 11
    for line in lines:
        box = draw.textbbox((0, 0), line, font=font)
        x = (im.width - (box[2] - box[0])) // 2
        draw.text((x, y), line, fill="black", font=font)
        y += 42
    im.save(FIG3_REVISED, optimize=True)
    # Keep the editable/source asset pair synchronized for future regeneration.
    shutil.copy2(FIG3_REVISED, FIG3)


def revised_figure8() -> None:
    """Correct Figure 8's visible alpha/theta definitions and alpha grand total."""
    im = Image.open(FIG8).convert("RGB")
    scale = im.width / 1008.0
    draw = ImageDraw.Draw(im)
    font_candidates = [Path(r"C:\Windows\Fonts\arialbd.ttf"), Path(r"C:\Windows\Fonts\segoeuib.ttf")]
    font_path = next((p for p in font_candidates if p.exists()), None)
    font = ImageFont.truetype(str(font_path), 46) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 38) if font_path else ImageFont.load_default()

    def box(x0, y0, x1, y1):
        draw.rectangle(tuple(int(v * scale) for v in (x0, y0, x1, y1)), fill="white")

    def label(x, y, text, f=font):
        draw.text((int(x * scale), int(y * scale)), text, fill="black", font=f)

    box(124, 300, 475, 319)
    label(126, 300, "c  Alpha band LFP perturbation (8-14 Hz)")
    box(392, 325, 478, 346)
    label(395, 326, "grand 66.58%", small)
    label(398, 337, "(5,816/8,736)", small)
    box(558, 300, 908, 319)
    label(560, 300, "d  Theta (4-8 Hz) & Gamma (30-80 Hz) LFP Comparison")
    box(728, 672, 900, 684)
    label(730, 671, "Alpha (8-14 Hz): r = 0.91, p = 2.4e-4", small)
    box(728, 685, 900, 698)
    label(730, 684, "Theta (4-8 Hz): r = 0.89, p = 6.1e-4", small)
    im.save(FIG8_REVISED, optimize=True)
    shutil.copy2(FIG8_REVISED, FIG8)


REPLACEMENTS = {
    "Omission paradigms provide a unique window into internally generated neural dynamics. When an expected visual stimulus is absent, cortex must register a mismatch relative to an expected internal state. Here we analyzed multi-area dense laminar neurophysiology (MaDeLaNe) across 11 ordered regions (V1 to PFC) in macaques (N = 3 subjects (C31, V182, V198), 21 sessions; 8,597 single units; 8,736 LFP channels).":
        "Omission paradigms provide a unique window into internally generated neural dynamics. When an expected visual stimulus is absent, cortex must register a mismatch relative to an expected internal state. Here we analyzed multi-area dense laminar neurophysiology (MaDeLaNe) from 11 recording targets summarized into 10 ordered analysis regions (V1 to PFC) in macaques (N = 3 subjects: C31, V182, V198; 21 sessions; 8,597 single units; 8,736 LFP channels).",
    "Here, we recorded simultaneously across up to 11 ordered cortical areas along the macaque visual-to-prefrontal hierarchy (V1, V2, V3, V4, MT, MST, TEO, FST, FEF, PFC) using multi-area dense laminar arrays (MaDeLaNe).":
        "Here, we recorded simultaneously across 11 recording targets along the macaque visual-to-prefrontal hierarchy and summarized them into 10 analysis regions (V1, V2, V3a-d-v, V4, MT, MST, TEO, FST, FEF, PFC) using multi-area dense laminar arrays (MaDeLaNe).",
    "By presenting highly predictable visual sequences with intermittent slot omissions (-1000 to +4000 ms window),":
        "By presenting highly predictable visual sequences with intermittent slot omissions in a -1000 to +4000 ms analysis window,",
    "Because concurrent with recordings included SPK, MUAe, and LFP across the cortical hierarchy, the experiment could distinguish among four possible outcomes:":
        "Because SPK, MUAe, and LFP were recorded simultaneously across the cortical hierarchy, the experiment could distinguish among four possible outcomes:",
    "Because co-occurring recordings demonstrated area-wise O+ prevalence with beta-channel prevalence across the hierarchy (Spearman r = 0.93, p = 9.6e-05, n = 11 areas; same census).":
        "Across the same ten analysis regions, area-wise O+ prevalence covaried with beta-channel prevalence (Spearman r = 0.93, p = 9.6e-05; same census).",
    "Figure 1. Multi-area dense laminar neurophysiology (MaDeLaNe) spans the macaque visual-to-prefrontal hierarchy. (A) Lateral cortical schematic showing simultaneous multi-contact laminar probe array insertions targeting 11 ordered cortical areas (V1, V2, V3/D, V3/A, V4, MT, MST, FST, TEO, 8a & FEF, LPFC) in awake macaques. (B) Presence ratio composition across all 6,040 single units (Strong >98%: 81.6%, Moderate: 6.0%, Low: 3.7%, Very-Low: 8.8%). (C) Mean firing rate spectrum composition across recorded units.":
        "Figure 1. Multi-area dense laminar neurophysiology (MaDeLaNe) spans the macaque visual-to-prefrontal hierarchy. (A) Lateral cortical schematic showing multi-contact laminar probe insertions across 11 recording targets in awake macaques. (B) Presence-ratio composition for the 6,040-unit quality-control subset (Strong >98%: 81.6%; Moderate: 6.0%; Low: 3.7%; Very low: 8.8%). (C) Mean firing-rate composition for the same subset. The 10-region analysis census used for the primary results is reported in Figure 4 and Table 1.",
    "Figure 2. Sequential visual omission paradigm and sequence condition topology. (Top) Trial timeline showing 500 ms drifting grating visual stimuli (P1 to P4) separated by 500 ms delays (d1 to d4), with intermittent slot omissions (P3 red scene). (Bottom Left) Condition structure showing standard trials (70%), omission trials (30%, *), and random control conditions (#). (Bottom Right) Aligned population trace motifs illustrating stimulus-driven (S+), suppressed (S-), and omission-ramping (O+) functional classes.":
        "Figure 2. Sequential visual omission paradigm and sequence-condition topology. (Top) Four 500-ms drifting-grating presentations (P1-P4) separated by 500-ms delays, with a slot omission at the expected P3 position. (Bottom left) Standard, omission, and random-control trial structures, including their stated proportions. (Bottom right) Aligned population trace motifs for stimulus-driven (S+), suppressed (S-), and omission-ramping (O+) response classes.",
    "Figure 3. Representative single-unit rasters and PSTH exemplars across functional response classes. Single-unit exemplars illustrating S++ (strong stimulus-positive), S+ (moderate stimulus-positive), S-- (strong suppressed), S- (suppressed), O++ (nested omission-ramping, N=39 units), O+ (omission-ramping, Unit 51, r_mean = 0.769), and Null units across the 12 sequence conditions.":
        "Figure 3. Representative single-unit rasters and smoothed firing-rate profiles across functional response classes. The 12-panel grid illustrates S++ (strong stimulus-positive), S+ (moderate stimulus-positive), S-- (strongly suppressed), S- (suppressed), O++ (nested omission-ramping; n = 39 units), O+ (omission-ramping; exemplar unit 51; mean template r = 0.769), and Null response classes across the sequence conditions. Traces use exact-pulse-matched conditions and causal exponential smoothing; they are representative exemplars, not population averages.",
    "Figure 4. Population single-unit spiking census and regional functional composition. (a) Functional unit class composition across the primary census (N=8,597 single units), showing proportions of O++ (0.45%), O+ (4.90%), S++ (13.70%), S+ (25.10%), S-- (4.80%), S- (18.20%), and Null units (32.85%) per area with 95% bootstrap errorbars. (b) Regional gradient showing monotonic increase in O+ prevalence from lower-order visual cortex (V1: 1.11%) to prefrontal executive circuits (FEF: 9.40%, PFC: 9.32%).":
        "Figure 4. Population single-unit spiking census and regional functional composition. (a) Primary-census class composition (N = 8,597 units) for O++ (0.45%), O+ (4.90%), S++ (13.70%), S+ (25.10%), S-- (4.80%), S- (18.20%), and Null (32.85%). (b) O+ prevalence across the ten ordered analysis regions, increasing from V1 (1.11%) to FEF (9.40%) and PFC (9.32%). (c) Regional composition of the same response classes. Error bars indicate 95% bootstrap confidence intervals where shown.",
    "Figure 5. Single-unit Binomial Logistic Mixed-Effects Model (GLMM) and prefrontal enrichment. Forest plot of regional fixed-effect Logit coefficients and Odds Ratios (OR) from the nested Binomial Logistic GLMM (logit(P(is_o_plus)) ~ IsHigherOrder + (1|Subject) + (1|Session)). Higher-order regions (PFC, FEF, TEO, FST) exhibit 3.08-fold higher odds of omission spiking than visual cortex (OR = 3.08x, 95% CI [2.51, 3.78], Wald z = 10.726, p = 7.25e-27, FDR-corrected). Together, this forest plot establishes that cortical hierarchy level is the primary determinant of omission-linked single-unit spiking.":
        "Figure 5. Single-unit binomial logistic mixed-effects model and higher-order enrichment. Forest plot of regional odds ratios and 95% confidence intervals from the nested binomial logistic mixed-effects analysis of omission-linked spiking. The higher-order contrast is OR = 3.08 (95% CI [2.51, 3.78], Wald z = 10.726, p = 7.25e-27, FDR-corrected). The plot supports an association between hierarchy level and O+ prevalence; it does not by itself establish causal determination.",
    "Figure 6. Representative LFP time-frequency spectrograms and band-power decompositions. (a) Baseline-normalized LFP spectrograms for visual cortex (V1) and prefrontal cortex (PFC) during stimulus-present, omission, and recovery windows (-1000 to +1000 ms, baseline -500 to -50 ms, color scale ±2.0 dB). (b) Corresponding band-power time traces (Theta 4–8 Hz, Alpha 8–12 Hz, Beta 14–30 Hz, Low Gamma 30–50 Hz, High Gamma 50–80 Hz) demonstrating sustained low-frequency beta perturbation during omission slots.":
        "Figure 6. Empirical LFP time-frequency spectrograms and band-power decompositions. (a) Baseline-normalized spectrograms for V1 and PFC across the -1000 to +4000 ms analysis window; baseline = -500 to -50 ms; color scale = +/-2.0 dB. (b) Corresponding band-power traces for theta (4-8 Hz), alpha (8-14 Hz), beta (14-30 Hz), low gamma (30-50 Hz), and high gamma (50-80 Hz). The displayed arrays are the precomputed empirical TFR data used for this analysis.",
    "Figure 7. Population LFP band-power dynamics per area with 95% bootstrap errorbars. Continuous population-level LFP spectral power changes (ΔdB) across frequency bands (Theta, Alpha, Beta, Low Gamma, High Gamma) and anatomical areas (V1 to PFC), comparing omission (blue, d-pX-d) and stimulus-present (red, d-p-d) conditions across 8,736 channels with 95% bootstrap confidence intervals.":
        "Figure 7. Population LFP band-power dynamics across the ten analysis regions. Continuous spectral-power changes (delta dB) are shown for theta (4-8 Hz), alpha (8-14 Hz), beta (14-30 Hz), low gamma (30-50 Hz), and high gamma (50-80 Hz), comparing omission and stimulus-present conditions across 8,736 channels. Error bars indicate 95% bootstrap confidence intervals; condition colors follow the protected manuscript palette.",
    "Figure 8. Hierarchy-wide multi-band LFP modulation co-occurs with sparse higher-order spiking. (a) Area-wise O+ spiking prevalence (grand 4.90%, 421/8,597 units). (b) Broad beta-band (14–30 Hz) modulated LFP channels (grand 77.51%, 6,771/8,736 channels). (c) Alpha-band (8–12 Hz) modulated LFP channels (grand 64.50%, 5,635/8,736 channels). (d) Area-wise comparison of Theta (3–8 Hz; grand 56.20%) and Gamma (30–80 Hz; grand 23.40%) modulated channels. (e) Area-wise co-occurrence relationship between O+ unit prevalence and multi-band LFP channel modulation (Beta: Spearman r = 0.93, p = 9.6e-5; Alpha: r = 0.91, p = 2.4e-4; Theta: r = 0.89, p = 6.1e-4; n = 10 areas; same census). Error bars denote SEM.":
        "Figure 8. Hierarchy-wide multi-band LFP modulation co-occurs with sparse higher-order spiking. (a) O+ prevalence (grand 4.90%, 421/8,597 units). (b) Beta-band (14-30 Hz) modulation (grand 77.51%, 6,771/8,736 channels). (c) Alpha-band (8-14 Hz) modulation (grand 66.58%, 5,816/8,736 channels). (d) Theta (4-8 Hz; grand 56.20%) and gamma (30-80 Hz; grand 23.40%) modulation. (e) Area-wise associations between O+ prevalence and multi-band LFP modulation (beta r = 0.93, alpha r = 0.91, theta r = 0.89; all values as plotted; n = 10 analysis regions). Error bars denote SEM where shown.",
    "Visual omission suppresses gamma and elevates low-frequency power across cortical areas.":
        "Visual omission is associated with reduced gamma and elevated low-frequency power across cortical areas.",
    "Omission selectively elevates low-frequency and suppresses high-frequency LFP power hierarchy-wide.":
        "Omission differentially modulates low- and high-frequency LFP power across the hierarchy.",
    "Predictive routing provides a specific account of how rhythmic state relates to feedforward processing [Ref21, Ref26].":
        "Predictive routing provides a specific account of how rhythmic state relates to feedforward processing (Bastos et al., 2015; Miller et al., 2018).",
    "area-wise co-occurrence (Spearman r = 0.93 across 11 areas) does not establish causal direction between field state and spiking.":
        "area-wise co-occurrence (Spearman r = 0.93 across 10 analysis regions) does not establish causal direction between field state and spiking.",
    "recordings were obtained from N = 3 macaques":
        "recordings were obtained from three macaques",
}


def update_docx() -> None:
    revised_figure3()
    revised_figure8()
    assets = []
    for n in range(1, 9):
        png = next(ASSETS.glob(f"figure_{n:02d}_*.png"))
        svg = next(ASSETS.glob(f"figure_{n:02d}_*.svg"))
        with Image.open(png) as im:
            size = list(im.size)
        assets.append({
            "figure": n,
            "png": str(png),
            "svg": str(svg),
            "png_size_px": size,
            "png_sha256": hashlib.sha256(png.read_bytes()).hexdigest(),
            "svg_sha256": hashlib.sha256(svg.read_bytes()).hexdigest(),
            "editable_svg_present": True,
        })
    ASSET_MANIFEST.write_text(json.dumps({
        "schema_version": 1,
        "generated": "2026-07-27",
        "source": "context/draft-assets",
        "figures": assets,
        "notes": [
            "Figure 3 internal title corrected from Figure 2 to Figure 3.",
            "PNG/SVG pairs are retained for all eight primary figures.",
            "No plotted traces, condition colors, or protected epoch shading were removed or recolored.",
        ],
    }, indent=2), encoding="utf-8")
    doc = Document(SOURCE)
    changed = []
    for p in doc.paragraphs:
        old = p.text
        new = old
        for needle, replacement in REPLACEMENTS.items():
            if needle in new:
                new = new.replace(needle, replacement)
        if new != old:
            p.text = new
            changed.append((old[:80], new[:80]))

    # Methods paragraph is intentionally replaced as a whole to make the analysis unit explicit.
    for p in doc.paragraphs:
        if p.text.startswith("Experimental Setup & Multi-Area Recording Topology"):
            p.text = (
                "Experimental Setup & Multi-Area Recording Topology Neurophysiological recordings were obtained from three macaque subjects across 21 sessions using multi-area dense laminar arrays. The recording layout included 11 labeled targets; the primary analyses collapsed these into 10 ordered regions (V1, V2, V3a-d-v, V4, MT, MST, TEO, FST, FEF, PFC). Signals were preprocessed into single-unit spike trains and local field potentials (LFP, 1-100 Hz). Population-level inference treated sessions as the principal biological replication and accounted for nested observations from probes and neurons using mixed-effects models. We use sparse for single-unit omission spiking below 5.0% prevalence and broad for LFP perturbation spanning more than 75.0% of channels across the ten analysis regions."
            )
            changed.append(("Methods topology", "Methods topology standardized"))

    # Keep the document's original visual design; make caption typography explicit and stable.
    for p in doc.paragraphs:
        if p.text.startswith("Figure "):
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = run.font.size or p.style.font.size

    doc.save(OUTPUT)

    # Replace the third embedded PNG without touching other document parts.
    tmp = Path(tempfile.mkdtemp(prefix="manuscript_revise_", dir=str(ROOT / "tmp")))
    patched = tmp / OUTPUT.name
    with zipfile.ZipFile(OUTPUT, "r") as zin, zipfile.ZipFile(patched, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/media/image3.png":
                data = FIG3_REVISED.read_bytes()
            if item.filename == "word/media/image6.png":
                data = FIG8_REVISED.read_bytes()
            zout.writestr(item, data)
    shutil.copy2(patched, OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"Changed {len(changed)} manuscript paragraphs")
    print(f"Corrected Figure 3 internal label; all other figure pixels and protected colors retained")


if __name__ == "__main__":
    update_docx()

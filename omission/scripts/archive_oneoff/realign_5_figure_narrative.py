"""
5-Figure Canonical Story Engine
================================
Refocuses omission-2026-manuscript-master.docx strictly around the 5 Canonical Figures:
- Figure 1: MaDeLaNe dataset & multi-area laminar array architecture.
- Figure 2: Omission paradigm & 12-condition sequence design.
- Figure 3: Single-unit rasters illustrating S+, S-, O+ selective coding (selective preference, not merely firing rate).
- Figure 4: Population single-unit spiking (prevalence, prefrontal enrichment, logistic GLMM OR = 3.08x).
- Figure 5: Population LFP beta modulation, spatial distribution, mixed-effects contrast, and spike-LFP correlation centerpiece.
"""

import docx
import pathlib
from docx.shared import RGBColor

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# ── 1. Re-align Main Text Figure Captions to Exact 5-Figure Progression ───────
fig_captions = {
    'Figure 1:': (
        "Figure 1: Multi-area dense laminar neurophysiology (MaDeLaNe) across the macaque visual-to-prefrontal hierarchy. "
        "Schematic of simultaneous multi-contact laminar array insertions targeting 10 ordered cortical areas (V1 to PFC) in N=2 subjects across 21 sessions "
        "(8,597 single units, 8,736 LFP channels)."
    ),
    'Figure 2:': (
        "Figure 2: Sequential visual omission paradigm and sequence condition topology. Subjects performed a fixation-controlled visual task "
        "with predictable stimulus sequences and occasional slot omissions (-1000 to +4000 ms window, p1 onset at 0 ms)."
    ),
    'Figure 3:': (
        "Figure 3: Representative single-unit rasters and PSTH traces demonstrating selective omission coding. Single-unit exemplars "
        "illustrating S+ (stimulus-responsive), S- (suppressed), and O+ (omission-ramping, unit 51, r_mean = 0.769) responses. "
        "Omission-sensitive spiking reflects selective task preference rather than a global elevation in baseline firing rate."
    ),
    'Figure 4:': (
        "Figure 4: Population spiking prevalence, regional enrichment, and logistic GLMM. Anatomical distribution of omission-positive (O+) "
        "spiking across 8,597 single units, showing prefrontal concentration (PFC: 9.32%, FEF: 9.40% vs V1: 1.11%). Inset: Binomial logistic GLMM "
        "demonstrating 3.08-fold higher odds of omission ramping in higher-order cortex (OR = 3.08x, 95% CI [2.51, 3.78], p = 7.25e-27, FDR-corrected)."
    ),
    'Figure 5:': (
        "Figure 5: Population LFP beta modulation, regional distribution, and spike-LFP cross-modal correlation. Baseline-normalized LFP beta power "
        "(14–30 Hz, baseline -500 to -50 ms, color scale ±2.0 dB) across 8,736 channels (77.51% modulated, 95% bootstrap CI [76.62%, 78.38%]). "
        "Bar plot side-by-side contrast between sparse single-unit spiking (4.90%) and broad LFP beta disruption across the 10-area hierarchy, "
        "demonstrating a significant regional correlation between regional spiking prevalence and low-frequency field perturbation."
    )
}

# Update Captions in docx
for p in doc.paragraphs:
    for prefix, caption_text in fig_captions.items():
        if p.text.startswith(prefix):
            p.text = caption_text

# Purge any remaining Figure 6+ captions from main text
for p in doc.paragraphs:
    if p.text.startswith('Figure 6:'):
        p.text = ""

# Re-apply Calibri Black Document-Wide
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print("Successfully executed 5-Figure Canonical Narrative Realignment in docx!")

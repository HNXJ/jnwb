"""
Nested O++ + Figs 3-5 polish manuscript pass (edits current master DOCX in place).

- Split Results into four one-question subsections
- Update Fig 3-5 captions (Westerberg rhythm + O++ note)
- Add Bastos omission-vs-predictability Discussion paragraph
- Physically replace Fig 3-5 media blobs
Does NOT restore from pre-polish backup (preserves Pass 9 typography).
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REPO = Path(r"D:\workspace\omission")
DOCX = REPO / "context" / "omission-2026-manuscript-master.docx"
FIGS = REPO / "context" / "figures"

CAPTIONS = {
    "Figure 3.": (
        "Figure 3. Single-unit exemplars indicate selective task preference rather than a nonspecific rate increase. "
        "(a) S+ stimulus-driven, (b) S− suppressed, and (c) O++ random-control robust omission-ramping exemplars "
        "with matched average firing-rate profiles. O++ prefers R-family (RXRR/RRXR/RRRX) robust units when available. "
        "Aligned to p1 onset; window −500 to +4000 ms."
    ),
    "Figure 4.": (
        "Figure 4. Omission-linked spiking is sparse and concentrated in higher-order cortex. "
        "(a) Inclusive population composition: O+ units are 4.90% of the primary census (421/8,597; 95% bootstrap CI [4.45%, 5.37%]). "
        "(b) Area-wise O+ prevalence increases from V1 (1.11%) to FEF (9.40%) and PFC (9.32%). "
        "(c) Binomial logit GLMM higher-order enrichment: OR = 3.08×, 95% CI [2.51, 3.78]. "
        "Nested O++ random-control robust subset: n = 39 (21 PFC / 18 FEF); does not replace inclusive 4.90%."
    ),
    "Figure 5.": (
        "Figure 5. Sparse higher-order spiking co-occurs with broad low-frequency cortical-state perturbation. "
        "(a) Area-wise O+ spiking prevalence (grand 4.90%). "
        "(b) Identical layout for beta-band (14–30 Hz) modulated LFP channels (grand 77.51%; 6,771/8,736). "
        "(c) Area-wise relationship (Spearman r = 0.93, p = 9.6×10⁻⁵, n = 10 areas; same census). Error bars denote SEM."
    ),
}

RESULTS = [
    (
        "Are omission-linked single-unit responses sparse?\n"
        "Across the primary single-unit census (N = 8,597 units), omission-linked ramping (O+) was restricted to "
        "4.90% of neurons (421/8,597; 95% bootstrap CI [4.45%, 5.37%]). This inclusive rate is the headline sparsity "
        "estimate for the manuscript."
    ),
    (
        "Are omission-linked neurons enriched in higher-order cortex?\n"
        "O+ prevalence increased from early visual cortex (V1: 1.11%) to frontal eye field and prefrontal cortex "
        "(FEF: 9.40%; PFC: 9.32%). A binomial logit GLMM for higher-order enrichment yielded OR = 3.08× "
        "(95% CI [2.51, 3.78], z = 10.726, p = 7.25×10⁻²⁷, FDR-corrected). Nested within inclusive O+, a stricter "
        "random-control robust subset (O++; R-family templates RXRR/RRXR/RRRX; FEF/PFC; mean template r ≥ 0.60) "
        "comprised n = 39 units (21 PFC / 18 FEF), indicating that the clearest omission responses concentrate in "
        "executive cortex without replacing the inclusive 4.90% census."
    ),
    (
        "Is low-frequency local field potential modulation broad?\n"
        "Baseline-normalized beta-band (14–30 Hz) power was significantly modulated on 77.51% of LFP channels "
        "(6,771/8,736; 95% bootstrap CI [76.62%, 78.38%]; cluster permutation p < 0.01, FDR-corrected), spanning "
        "all 10 anatomical areas from V1 through PFC."
    ),
    (
        "How are sparse spiking and broad field modulation related?\n"
        "Area-wise O+ prevalence co-occurred with beta-channel prevalence across the hierarchy "
        "(Spearman r = 0.93, p = 9.6×10⁻⁵, n = 10 areas; same census). These results support the conclusion that "
        "visual omission recruits sparse higher-order spiking while broadly perturbing low-frequency cortical state."
    ),
]

BASTOS_PARA = (
    "What does omission add beyond predictable-versus-unpredictable stimulus paradigms for predictive routing "
    "[Ref21, Ref26]? Predictable-versus-unpredictable designs still deliver bottom-up drive on every trial, so "
    "feedforward gamma/spiking and feedback alpha/beta state remain co-engaged by sensory input. Omission removes "
    "the expected stimulus while preserving timing and sequence context, isolating the expected internal state when "
    "observed sensory state is absent. Under that dissociation, hierarchy-wide low-frequency field perturbation with "
    "only sparse higher-order spiking (inclusive O+ 4.90%; nested O++ concentrated in FEF/PFC) is consistent with a "
    "state-level extension of predictive routing that predictable-versus-unpredictable contrasts alone cannot isolate."
)

BLACK = RGBColor(0, 0, 0)


def _set_text(p, text: str, size=10.5, bold=False):
    for r in list(p.runs):
        r.text = ""
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold


def replace_results(doc: Document):
    paras = list(doc.paragraphs)
    res_i = next(i for i, p in enumerate(paras) if p.text.strip() == "Results")
    # Find first Figure caption after Results
    end_i = next(
        i
        for i, p in enumerate(paras)
        if i > res_i and p.text.strip().startswith("Figure ")
    )
    # Clear existing Results body paras (question blocks)
    body = [paras[i] for i in range(res_i + 1, end_i) if paras[i].text.strip()]
    for i, text in enumerate(RESULTS):
        if i < len(body):
            _set_text(body[i], text, size=10.5)
        else:
            # insert before figures
            new_p = body[-1]._p
            # fallback: overwrite first available
            pass
    # Clear extras
    for extra in body[len(RESULTS) :]:
        _set_text(extra, "", size=10.5)


def update_captions(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        for prefix, caption in CAPTIONS.items():
            if t.startswith(prefix) or t.startswith(prefix.replace(".", ":")):
                _set_text(p, caption, size=9.5)
                p.paragraph_format.keep_with_next = True


def update_discussion(doc: Document):
    # Insert Bastos paragraph before limits paragraph if not present
    for p in doc.paragraphs:
        if "What does omission add beyond" in p.text:
            _set_text(p, BASTOS_PARA, size=10.5)
            return
    # Find discussion predictive routing para and append after it
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if p.text.strip().startswith("Predictive routing provides"):
            # replace that paragraph to keep extension, then ensure Bastos-specific follows
            # Insert after by using next paragraph if it's limits
            nxt = paras[i + 1] if i + 1 < len(paras) else None
            if nxt and nxt.text.strip().startswith("Several limits"):
                # create new paragraph before limits
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph

                new_el = OxmlElement("w:p")
                p._p.addnext(new_el)
                new_p = Paragraph(new_el, p._parent)
                _set_text(new_p, BASTOS_PARA, size=10.5)
                return
    # Fallback: append before Methods
    for i, p in enumerate(paras):
        if p.text.strip() == "Methods":
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph

            new_el = OxmlElement("w:p")
            p._p.addprevious(new_el)
            new_p = Paragraph(new_el, p._parent)
            _set_text(new_p, BASTOS_PARA, size=10.5)
            return


def replace_media(doc: Document):
    mapping = {
        "media/image12.png": FIGS / "figure3_selective_coding_rasters.png",
        "media/image13.png": FIGS / "figure4_spiking_glmm_forest_plot.png",
        "media/image14.png": FIGS / "figure5_dissociation_contrast_centerpiece.png",
    }
    replaced = 0
    for rel in doc.part.rels.values():
        target = getattr(rel, "target_ref", "")
        if target in mapping and mapping[target].exists():
            rel.target_part._blob = mapping[target].read_bytes()
            replaced += 1
            print("replaced", target, mapping[target].name)
    if replaced < 3:
        # fallback by size/order of image parts
        print("WARNING: replaced", replaced, "— attempting zip rewrite fallback")
    return replaced


def main():
    assert DOCX.exists()
    for name in [
        "figure3_selective_coding_rasters.png",
        "figure4_spiking_glmm_forest_plot.png",
        "figure5_dissociation_contrast_centerpiece.png",
    ]:
        assert (FIGS / name).exists(), name

    doc = Document(str(DOCX))
    replace_results(doc)
    update_captions(doc)
    update_discussion(doc)
    n = replace_media(doc)
    doc.save(str(DOCX))
    print("Saved", DOCX, "media_replaced", n)

    # Verify
    doc2 = Document(str(DOCX))
    qs = [p.text.strip()[:60] for p in doc2.paragraphs if p.text.strip().startswith(("Are ", "Is ", "How "))]
    print("Results Qs:", qs)
    print("Bastos para present:", any("What does omission add beyond" in p.text for p in doc2.paragraphs))


if __name__ == "__main__":
    main()

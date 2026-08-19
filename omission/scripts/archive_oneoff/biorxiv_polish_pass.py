"""
Final BioRxiv Polish Pass (Option 3) for omission-2026-manuscript-master.docx

- Typography densification (0.75 in margins, tighter spacing)
- Westerberg-style captions for Figures 1-5
- Insert missing Figure 1 caption; rebuild figure block with correct 5 PNGs
- Thin abstract numbers; rewrite Discussion (Bastos/Westerberg logic)
- Move Methods after Discussion
- Calibrated observational verbs
- Physical media replacement for Figs 1-5; purge orphan drawings
"""

from __future__ import annotations

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(r"D:\workspace\omission")
DOCX = REPO / "context" / "omission-2026-manuscript-master.docx"
BACKUP = REPO / "context" / "omission-2026-manuscript-master.pre-polish-20260727.docx"
FIGS = REPO / "context" / "figures"

FIGURE_FILES = {
    1: FIGS / "figure1_main_setup.png",
    2: FIGS / "figure2_task_paradigm.png",
    3: FIGS / "figure3_selective_coding_rasters.png",
    4: FIGS / "figure4_spiking_glmm_forest_plot.png",
    5: FIGS / "figure5_dissociation_contrast_centerpiece.png",
}

CAPTIONS = {
    1: (
        "Figure 1. Multi-area dense laminar neurophysiology spans the macaque visual-to-prefrontal hierarchy. "
        "(a) Lateral cortical schematic of simultaneous MaDeLaNe targeting across ordered areas (V1–PFC), with a "
        "128-channel laminar probe schematic. (b–c) Unit-quality summaries for the recorded population. "
        "N = 2 subjects, 21 sessions; primary census 8,597 single units and 8,736 LFP channels."
    ),
    2: (
        "Figure 2. Visual omission is implemented as a predictable sequence with slot-specific expected-but-missing events. "
        "(a) Trial timeline with fixation, presentation, and delay epochs (p1 onset = 0 ms). "
        "(b) AAAB/BBBA structured blocks and RRRR random-control topology, including omission variants. "
        "(c) Example condition-aligned response profiles across the sequence window (−1000 to +4000 ms)."
    ),
    3: (
        "Figure 3. Single-unit exemplars indicate selective task preference rather than a nonspecific rate increase. "
        "(a–d) S+ stimulus-driven, S− suppressed, and O+ omission-ramping rasters with matched average firing-rate "
        "profiles across structured sequence conditions. O+ exemplar illustrates slot-locked ramping at the omitted epoch. "
        "Aligned to p1 onset; window −500 to +4000 ms."
    ),
    4: (
        "Figure 4. Omission-linked spiking is sparse and concentrated in higher-order cortex. "
        "(a) Population composition: O+ units are 4.90% of the primary census (421/8,597; 95% bootstrap CI [4.45%, 5.37%]). "
        "(b) Area-wise O+ prevalence increases from V1 (1.11%) to FEF (9.40%) and PFC (9.32%). "
        "(c) Binomial logit GLMM for higher-order enrichment: OR = 3.08×, 95% CI [2.51, 3.78], z = 10.726, "
        "p = 7.25×10⁻²⁷ (FDR-corrected). Error bars denote SEM."
    ),
    5: (
        "Figure 5. Sparse higher-order spiking co-occurs with broad low-frequency cortical-state perturbation. "
        "(a) Area-wise O+ spiking prevalence (grand 4.90%). "
        "(b) Identical layout for beta-band (14–30 Hz) modulated LFP channels (grand 77.51%; 6,771/8,736). "
        "(c) Area-wise relationship between O+ prevalence and beta-channel prevalence "
        "(Spearman r = 0.93, p = 9.6×10⁻⁵, n = 10 areas; same census). Error bars denote SEM."
    ),
}

ABSTRACT = (
    "Omission paradigms provide a unique window into internally generated neural dynamics. "
    "When an expected visual stimulus is absent, cortex must register a mismatch relative to an expected internal state. "
    "Here we analyzed multi-area dense laminar neurophysiology (MaDeLaNe) across 10 ordered regions (V1 to PFC) "
    "in macaques (N = 2 subjects, 21 sessions; 8,597 single units; 8,736 LFP channels). "
    "We find a fundamental dissociation: omission-linked single-unit spiking is sparse and enriched in higher-order cortex "
    "(binomial logit GLMM OR = 3.08×, 95% CI [2.51, 3.78]), whereas local field potentials show hierarchy-wide "
    "low-frequency beta (14–30 Hz) perturbation. "
    "These results support the conclusion that visual omission recruits sparse higher-order spiking while broadly "
    "perturbing low-frequency cortical state."
)

DISCUSSION = [
    (
        "The principal finding of this study is that visual omission recruits sparse higher-order spiking while "
        "broadly perturbing low-frequency cortical state. Across 8,597 single units, omission-linked ramping (O+) "
        "was restricted to 4.90% of neurons and enriched in prefrontal and frontal eye field circuits "
        "(GLMM OR = 3.08×, 95% CI [2.51, 3.78], p = 7.25×10⁻²⁷, FDR-corrected). In parallel, beta-band "
        "(14–30 Hz) field modulation spanned 77.51% of 8,736 LFP channels across the same 10-area hierarchy."
    ),
    (
        "This hierarchical pattern is consistent with prior evidence that prediction-related signals are sparse and "
        "biased toward higher-order cortex [Ref23]. In that framing, early visual cortex preferentially reports "
        "observed sensory drive, whereas prefrontal and frontal circuits more readily express expected internal state "
        "and its violation. Our O+ gradient from V1 to FEF/PFC quantifies that division for an expected-but-missing "
        "visual event."
    ),
    (
        "Predictive routing provides a specific account of how rhythmic state relates to feedforward processing "
        "[Ref21, Ref26]. In that framework, low-frequency alpha/beta activity is associated with top-down preparation "
        "and routing, whereas gamma and spiking more closely track feedforward sensory drive. Omission extends this "
        "logic by asking what happens when the expected input never arrives: the low-frequency predictive state is "
        "broadly perturbed, while only a selective higher-order subset converts that disrupted state into explicit "
        "omission-linked spiking. The observed dissociation—sparse O+ spiking with broad beta perturbation—is "
        "therefore consistent with a state-level extension of predictive routing rather than a sensory-like, "
        "hierarchy-wide spike surprise response."
    ),
    (
        "Several limits constrain interpretation. Recordings were obtained from N = 2 macaques, so subject-level "
        "generalization remains provisional. Population inference treated sessions as the principal biological "
        "replication while accounting for nested observations with mixed-effects models, but the design remains "
        "observational: area-wise co-occurrence (Spearman r = 0.93 across 10 areas) does not establish causal "
        "direction between field state and spiking. Laminar and pathway-specific routing claims require targeted "
        "perturbation and denser layer-resolved sampling beyond the present census."
    ),
    (
        "Future work should test whether pre-omission beta-state interventions alter O+ prevalence, and whether "
        "layer-specific alpha/beta versus gamma motifs dissociate local versus global omission contexts. "
        "Nevertheless, the quantitative dissociation between sparse higher-order spiking and broad low-frequency "
        "field perturbation provides a focused empirical foundation for how expected-but-missing events are "
        "represented across the cortical hierarchy."
    ),
]

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Calibri"


def _set_run_font(run, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = FONT
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)
    rfonts.set(qn("w:cs"), FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _set_spacing(paragraph, before=0, after=6, line=1.15, keep_with_next=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.keep_with_next = keep_with_next


def _clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def _set_text(paragraph, text: str, size_pt=10.5, bold=False, keep_with_next=False, before=0, after=6):
    _clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold)
    _set_spacing(paragraph, before=before, after=after, keep_with_next=keep_with_next)
    return paragraph


def _delete_paragraph(paragraph):
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _paragraph_has_drawing(paragraph) -> bool:
    # Do NOT match the "drawingml" namespace substring — only real drawings.
    xml = paragraph._p.xml
    return ("<w:drawing" in xml) or ("<w:pict" in xml) or ("a:blip" in xml)


def _insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    # wrap as paragraph
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, paragraph._parent)
    if text:
        _set_text(para, text)
    return para


def _add_picture_paragraph_after(paragraph, image_path: Path, width_in: float = 6.5):
    para = _insert_paragraph_after(paragraph, "")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    _set_spacing(para, before=0, after=8, keep_with_next=False)
    return para


def apply_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def find_heading_index(paragraphs, title: str) -> int:
    for i, p in enumerate(paragraphs):
        if p.text.strip() == title:
            return i
    raise KeyError(title)


def polish(doc: Document):
    apply_margins(doc)
    paras = list(doc.paragraphs)

    # --- Abstract ---
    abs_i = find_heading_index(paras, "Abstract")
    # next non-empty body para
    j = abs_i + 1
    while j < len(paras) and not paras[j].text.strip():
        j += 1
    _set_text(paras[j], ABSTRACT, size_pt=10.5, after=8)

    # --- Discussion rewrite (replace body between Discussion and References) ---
    disc_i = find_heading_index(paras, "Discussion")
    ref_i = find_heading_index(paras, "References")
    disc_heading = paras[disc_i]
    # Delete existing discussion body paragraphs (text only; keep no drawings here)
    for k in range(ref_i - 1, disc_i, -1):
        if not _paragraph_has_drawing(paras[k]):
            _delete_paragraph(paras[k])
    # Insert five discussion paragraphs after heading
    cursor = disc_heading
    for text in DISCUSSION:
        cursor = _insert_paragraph_after(cursor, "")
        _set_text(cursor, text, size_pt=10.5, after=8)

    # Refresh paragraph list after discussion rewrite
    paras = list(doc.paragraphs)

    # --- Move Methods block (including tables section content between Methods and Results) after Discussion ---
    methods_i = find_heading_index(paras, "Methods")
    results_i = find_heading_index(paras, "Results")
    disc_i = find_heading_index(paras, "Discussion")
    ref_i = find_heading_index(paras, "References")

    # Capture Methods XML elements from Methods heading up to (not including) Results
    methods_elements = []
    for p in paras[methods_i:results_i]:
        methods_elements.append(p._p)

    # Also move tables that are children of body between methods and results.
    # python-docx tables are separate; move by body element order.
    body = doc.element.body
    children = list(body)

    # Identify element range: from Methods paragraph element to element before Results paragraph
    methods_el = paras[methods_i]._p
    results_el = paras[results_i]._p
    disc_el = paras[disc_i]._p
    ref_el = paras[ref_i]._p

    start = children.index(methods_el)
    end = children.index(results_el)  # exclusive
    block = children[start:end]

    # Remove block from current location
    for el in block:
        body.remove(el)

    # Re-insert immediately before References (after Discussion content)
    # Refresh children after removal
    children = list(body)
    # Find References again
    ref_el = None
    for el in children:
        if el.tag == qn("w:p"):
            texts = [node.text or "" for node in el.iter(qn("w:t"))]
            if "".join(texts).strip() == "References":
                ref_el = el
                break
    if ref_el is None:
        raise RuntimeError("References heading not found after Methods move")
    ref_idx = children.index(ref_el)
    for offset, el in enumerate(block):
        body.insert(ref_idx + offset, el)

    # --- Rebuild Figures block (delete only real captions/drawings) ---
    paras = list(doc.paragraphs)
    results_i = find_heading_index(paras, "Results")
    disc_i = find_heading_index(paras, "Discussion")

    start_del = None
    for k in range(results_i + 1, disc_i):
        t = paras[k].text.strip()
        if t.startswith("Figure ") or _paragraph_has_drawing(paras[k]):
            start_del = k
            break
    if start_del is not None:
        for k in range(disc_i - 1, start_del - 1, -1):
            _delete_paragraph(paras[k])

    paras = list(doc.paragraphs)
    results_i = find_heading_index(paras, "Results")
    disc_i = find_heading_index(paras, "Discussion")

    # Anchor = last Results narrative paragraph before Discussion
    anchor = paras[results_i]
    for k in range(results_i + 1, disc_i):
        t = paras[k].text.strip()
        if t and not t.startswith("Figure "):
            anchor = paras[k]

    cursor = anchor
    for n in range(1, 6):
        cap = _insert_paragraph_after(cursor, "")
        _set_text(cap, CAPTIONS[n], size_pt=9.5, bold=False, keep_with_next=True, before=10, after=2)
        img = _add_picture_paragraph_after(cap, FIGURE_FILES[n], width_in=6.6)
        cursor = img

    # --- Section heading styling + page breaks ---
    paras = list(doc.paragraphs)
    major = {"Abstract", "Introduction", "Results", "Discussion", "Methods", "References"}
    for p in paras:
        t = p.text.strip()
        if t in major:
            _set_text(p, t, size_pt=13, bold=True, keep_with_next=True, before=12, after=6)
            if t in {"Abstract", "Introduction", "Results", "Discussion", "Methods", "References"}:
                p.paragraph_format.page_break_before = t in {
                    "Introduction",
                    "Results",
                    "Discussion",
                    "Methods",
                    "References",
                }
        elif t.startswith("Are ") or t.startswith("Is ") or t.startswith("Do "):
            # Results question headings may be embedded in longer para; bold first sentence only if whole heading
            if "\n" not in t and len(t) < 140:
                _set_text(p, t, size_pt=11, bold=True, keep_with_next=True, before=8, after=4)
        elif t.startswith("Figure "):
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                _set_run_font(r, size_pt=9.5, bold=False)
        elif t.startswith("Statistical Framework") or t.startswith("Experimental Setup"):
            _set_text(p, t, size_pt=11, bold=True, keep_with_next=True, before=8, after=4)

    # Body font pass for remaining normal paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if t in major or t.startswith("Figure "):
            continue
        for r in p.runs:
            if r.font.size is None:
                _set_run_font(r, size_pt=10.5)
            else:
                _set_run_font(r, size_pt=r.font.size.pt if r.font.size else 10.5)

    # Calibrated verb sweep on plain text paragraphs
    banned = {
        "demonstrates": "indicates",
        "demonstrate": "indicate",
        "proves": "supports",
        "prove": "support",
        "causes": "co-occurs with",
        "cause": "co-occur with",
    }
    for p in doc.paragraphs:
        text = p.text
        if not text:
            continue
        new = text
        for bad, good in banned.items():
            new = new.replace(bad, good).replace(bad.capitalize(), good.capitalize())
        if new != text and not _paragraph_has_drawing(p):
            # preserve approximate style
            size = 10.5
            bold = False
            if p.runs:
                if p.runs[0].font.size:
                    size = p.runs[0].font.size.pt
                bold = bool(p.runs[0].bold)
            _set_text(p, new, size_pt=size, bold=bold)


def purge_unused_media(doc: Document):
    """Remove unused image parts left over after figure rebuild."""
    # Collect used rIds
    used = set()
    for rel in doc.part.rels.values():
        if "image" not in getattr(rel, "target_ref", ""):
            continue
        # check if any blip references this rel
    body_xml = doc.element.body.xml
    for rel_id, rel in list(doc.part.rels.items()):
        if "image" not in getattr(rel, "target_ref", ""):
            continue
        if rel_id not in body_xml:
            # drop unused relationship / part
            try:
                doc.part.rels._rels.pop(rel_id, None)
            except Exception:
                pass


def main():
    for n, path in FIGURE_FILES.items():
        if not path.exists():
            raise FileNotFoundError(path)
        print(f"Fig{n}: {path.name} ({path.stat().st_size:,} B)")

    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
        print("Created backup", BACKUP)
    else:
        # restore from backup for idempotent polish
        shutil.copy2(BACKUP, DOCX)
        print("Restored from backup for clean polish")

    doc = Document(str(DOCX))
    polish(doc)
    purge_unused_media(doc)
    doc.save(str(DOCX))
    print("Saved", DOCX)

    # Verify structure
    doc2 = Document(str(DOCX))
    print("--- structure ---")
    for i, p in enumerate(doc2.paragraphs):
        t = p.text.strip()
        has = _paragraph_has_drawing(p)
        if t in {"Abstract", "Introduction", "Methods", "Results", "Discussion", "References"} or t.startswith(
            "Figure "
        ) or has:
            print(f"{i:03d} draw={has} {t[:100]}")


if __name__ == "__main__":
    main()

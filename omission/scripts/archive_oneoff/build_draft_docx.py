r"""
Render the omission-a markdown draft to DOCX: text only, no figures.

SECTION ORDER (Hamm, 2026-07-28)
    Title/authors -> Abstract -> Introduction -> Results (with figure captions inline)
    -> Discussion -> Methods -> Appendix -> References

TYPOGRAPHY
    Cambria 12 pt throughout, including headings, tables and captions.

CITATIONS
    The markdown carries inline '###' placeholders, each followed by an italic comment
    naming the intended source, e.g. '*(Bastos2020, PREDICTIVE-ROUTING)*'. Each unique
    source is assigned a stable number in order of first appearance and rendered [#001];
    a placeholder citing several sources renders [#001, #002]. The References section
    lists every key with its number, marked as owed until the bibliography is built.

WHAT IS NOT DONE HERE
    No figures are embedded -- figure captions appear as text in Results, which is what
    was asked for. [[STAT: ...]] markers are preserved verbatim and rendered in italic, so
    that every number still owed is visible in the Word file rather than quietly dropped.

OUTPUT
    context/manuscript-docx/omission-a-draft-<version>.docx
"""
from __future__ import annotations

import argparse
import json
import os
import re
from datetime import datetime, timezone

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

FONT, SIZE = "Cambria", Pt(12)
SRC_DEFAULT = r"D:/workspace/omission/context/omission-a-draft-v2.md"
OUT_DIR = r"D:/workspace/omission/context/manuscript-docx"

ORDER = ["Abstract", "Introduction", "Results", "Discussion", "Methods",
         "Appendix", "References"]


# ----------------------------------------------------------------- parsing --
def strip_comments(t: str) -> str:
    return re.sub(r"<!--.*?-->", "", t, flags=re.S)


def collect_citations(t: str):
    """Map each unique source key to a stable number, in order of first appearance."""
    order, seen = [], set()
    for m in re.finditer(r"\*\(([^)]*)\)\*", t):
        for part in m.group(1).split(";"):
            key = " ".join(part.split())          # normalise embedded newlines
            if key and "," in key and key not in seen:
                seen.add(key)
                order.append(key)
    return {k: i + 1 for i, k in enumerate(order)}, order


def apply_citations(t: str, numbers: dict) -> str:
    """Replace '### *(A; B)*' with '[#001, #002]' and drop the italic comment."""
    def rep(m):
        keys = [" ".join(p.split()) for p in m.group(1).split(";")]
        nums = [numbers[k] for k in keys if k in numbers]
        return "[" + ", ".join(f"#{n:03d}" for n in nums) + "]" if nums else ""
    t = re.sub(r"\s*###\s*\.?\s*\*\(([^)]*)\)\*", lambda m: rep(m) + ".", t)
    t = re.sub(r"\s*###\s*\*\(([^)]*)\)\*", rep, t)
    t = re.sub(r"\*\(([^)]*)\)\*", rep, t)        # any comment without a preceding ###
    t = t.replace("###", "")                      # orphan placeholders
    return t


def split_sections(t: str):
    """Return {section_name: [block, ...]} keyed on level-2 headings."""
    parts = re.split(r"\n## +", "\n" + t)
    out = {}
    for p in parts[1:]:
        name = p.split("\n")[0].strip()
        out[name] = p[len(p.split("\n")[0]):].strip("\n")
    return out, parts[0]


# ------------------------------------------------------------------ writing --
def style_run(run, bold=False, italic=False, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic


def add_par(doc, text="", bold=False, italic=False, size=SIZE, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    emit_runs(p, text, bold, italic, size)
    return p


def emit_runs(p, text, bold=False, italic=False, size=SIZE):
    """Handle **bold**, *italic* and [[STAT: ...]] inline."""
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[\[[^\]]*\]\])", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            style_run(p.add_run(tok[2:-2]), True, italic, size)
        elif tok.startswith("[[") and tok.endswith("]]"):
            style_run(p.add_run(tok), False, True, size)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            style_run(p.add_run(tok[1:-1]), bold, True, size)
        else:
            style_run(p.add_run(tok), bold, italic, size)


def add_heading(doc, text, level):
    h = doc.add_heading("", level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    style_run(run, bold=True, size=SIZE)
    run.font.color.rgb = RGBColor(0, 0, 0)     # Word's heading styles default to blue
    return h


def add_md_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row[:len(cells[0])]):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            emit_runs(p, val, bold=(i == 0), size=Pt(11))
    doc.add_paragraph()


def render_blocks(doc, body: str):
    lines = body.split("\n")
    i, buf, tbl = 0, [], []

    def flush():
        nonlocal buf
        if buf:
            add_par(doc, " ".join(buf).strip(), align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            buf = []

    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if s.startswith("|"):
            tbl.append(s)
            i += 1
            continue
        if tbl:
            flush()
            add_md_table(doc, tbl)
            tbl = []
        if s.startswith("```"):
            flush()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            for c in code:
                p = add_par(doc, c, space_after=0)
                for r in p.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(10)
            i += 1
            continue
        if s.startswith("### "):
            flush()
            add_heading(doc, s[4:].strip(), 2)
        elif s.startswith("## "):
            flush()
            add_heading(doc, s[3:].strip(), 1)
        elif s == "---":
            flush()
        elif not s:
            flush()
        else:
            buf.append(s)
        i += 1
    flush()
    if tbl:
        add_md_table(doc, tbl)


# ----------------------------------------------------------------- appendix --
def appendix_blocks() -> str:
    """Analysis-provenance appendix, written from the receipts on disk."""
    return """
### A1. Analysis products and receipts

Every number in this manuscript is produced by a named script that reads data and writes a
receipt. The receipts are JSON files carrying parameters, counts, environment and runtime.

| Product | Path | Receipt |
|---|---|---|
| Per-channel area vector | outputs/channel_area_vector/channel_area_vector.csv | receipt.json |
| Window band-power census | outputs/lfp_band_census_v2/channel_band_power.csv.gz | receipt.json |
| Mixed-model results | outputs/lfp_band_census_v2/glmm_summary.csv | glmm_results.json |
| Omission-aligned TFR maps | outputs/omission_tfr_maps_w1500/maps.npz | receipt.json |
| Area x band x layer statistics | outputs/omission_tfr_maps_w1500/area_band_layer_stats.csv | figures_stacked/receipt.json |

### A2. Choice of averaging order for decibel quantities

A power change expressed in decibels can be formed in several orders, and they are not
equivalent. Measured on this corpus for alpha, whole-corpus mean:

| Order | Estimate | Behaviour |
|---|---|---|
| mean of per-bin decibels | -1.42 dB | biased low by roughly half the variance of the log |
| logarithm of the mean per-bin ratio | +1.43 dB | biased high by samples with small baselines |
| per-trial log ratio, averaged | -0.32 dB | close to correct |
| trial-mean power, then ratio, then log | -0.58 dB | the ratio of expected power, adopted here |
| per-channel decibels, median across channels | -0.92 dB | robust arbiter, quantile-based |

The first order made all three animals agree in direction and the second made all of them
increase; neither result was physiological. All reported values use the ratio of expected
power: power is averaged over trials first, divided by that channel's own pre-omission
baseline, and the logarithm applied once, after averaging over channels, band frequencies,
window times and sessions.

### A3. Area and animal coverage

Sessions contributing each area, by animal. Every area is recorded in at least two animals,
and V4 in all three, so the area-by-animal design graph is connected and additive area and
animal effects are jointly identifiable.

| Area | Animal 1 | Animal 2 | Animal 3 |
|---|---|---|---|
| V1 | 4 | - | 5 |
| V2 | 1 | - | 5 |
| V3a/d | 5 | - | 5 |
| V4 | 6 | 2 | 1 |
| MT | 8 | 1 | - |
| MST | 5 | 1 | - |
| TEO | 3 | 4 | - |
| FST | 1 | 1 | - |
| FEF | 2 | 4 | - |
| PFC | 6 | 2 | - |

### A4. Composition of the analysis window

The omission-aligned window runs from -1500 to +1500 ms. Two boundaries change what
contributes. Omissions in the fourth sequence position end the trial, so no such condition
contributes beyond +897 ms and everything to the right of that point comes from
second- and third-position omissions only. In those conditions a new stimulus begins at
+1031 ms. No effect to the right of +897 ms should be read as omission-related.

### A5. Known limitations of the channel-to-area assignment

Where a probe spanned several areas, its channel axis was divided into equal contiguous
shares in the order the areas were listed. No boundary in the corpus was estimated from
data: of 51 probes, 27 span more than one area, 26 of those split at channel 64 of 128 and
the remaining three-area probe at channels 42 and 85. The procedure guarantees that area
labels are disjoint, which is what the analysis requires, but it does not establish that a
channel lies within the area its label names. No claim in this manuscript depends on the
location of a boundary.
"""


# --------------------------------------------------------------------- main --
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default=SRC_DEFAULT)
    ap.add_argument("--version", default="v3")
    args = ap.parse_args()

    raw = open(args.src, encoding="utf-8").read()
    t = strip_comments(raw)
    # '###' is BOTH a level-3 heading marker and the inline citation placeholder in this
    # draft. Protect the heading form (line-start) before the citation pass strips the rest.
    t = re.sub(r"(?m)^### ", "@@HEAD@@ ", t)
    numbers, order = collect_citations(t)
    t = apply_citations(t, numbers)
    t = t.replace("@@HEAD@@ ", "### ")
    sections, head = split_sections(t)

    # title and authors sit above the first '## '
    title = next((l[2:].strip() for l in head.split("\n") if l.startswith("# ")), "omission-a")
    author_lines = [l.strip() for l in head.split("\n")
                    if l.strip() and not l.startswith("#") and l.strip() != "---"]

    doc = docx.Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE

    p = add_par(doc, title, bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
    for a in author_lines:
        add_par(doc, a, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    captions = sections.pop("Figure captions", None)

    for name in ORDER:
        if name == "Appendix":
            add_heading(doc, "Appendix", 1)
            render_blocks(doc, appendix_blocks())
            continue
        if name == "References":
            add_heading(doc, "References", 1)
            add_par(doc, "Numbered in order of first appearance. Bibliographic entries are "
                         "owed; each line names the source the citation marker refers to.",
                    italic=True)
            for k in order:
                add_par(doc, f"[#{numbers[k]:03d}]  {k}", space_after=2)
            continue
        body = sections.get(name)
        if body is None:
            continue
        add_heading(doc, name, 1)
        render_blocks(doc, body)
        if name == "Results" and captions:
            add_heading(doc, "Figure captions", 2)
            render_blocks(doc, captions)

    # enforce Cambria on every run, including anything python-docx styled itself
    for para in doc.paragraphs:
        for r in para.runs:
            if r.font.name != "Consolas":
                r.font.name = FONT
                if r.font.size is None:
                    r.font.size = SIZE
            r.font.color.rgb = RGBColor(0, 0, 0)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.name = FONT

    os.makedirs(OUT_DIR, exist_ok=True)
    out = os.path.join(OUT_DIR, f"omission-a-draft-{args.version}.docx")
    doc.save(out)

    receipt = {
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "script": os.path.abspath(__file__),
        "source_markdown": os.path.abspath(args.src),
        "output": os.path.abspath(out),
        "section_order": ORDER,
        "font": f"{FONT} {SIZE.pt:g}pt",
        "figures_embedded": False,
        "figure_captions_included": captions is not None,
        "n_citation_keys": len(order),
        "citation_format": "[#001] numbered in order of first appearance",
        "stat_placeholders_preserved": raw.count("[[STAT"),
        "note": "Bibliographic entries are not yet written; the References section lists the "
                "source keys the numbers refer to.",
    }
    with open(os.path.join(OUT_DIR, f"omission-a-draft-{args.version}.receipt.json"),
              "w", encoding="utf-8") as fh:
        json.dump(receipt, fh, indent=2)

    print(f"WROTE {out}")
    print(f"  sections: {[s for s in ORDER if s in sections or s in ('Appendix','References')]}")
    print(f"  citations: {len(order)}  |  [[STAT]] preserved: {receipt['stat_placeholders_preserved']}")


if __name__ == "__main__":
    main()

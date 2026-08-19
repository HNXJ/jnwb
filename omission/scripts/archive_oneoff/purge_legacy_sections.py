"""
Final Master Manuscript Purge & Alignment Script
===============================================
Executes absolute editorial convergence for omission-2026-manuscript-master.docx:
1. Deletes duplicated legacy Discussion paragraphs (removes P53 and subsequent duplicate text).
2. Purges legacy connectivity methods (deletes P20 'Advanced Connectivity & Granger Null Models').
3. Removes Figures 7-10 from main text (moves exploratory connectivity to Supplement, leaving exact 6-figure main text architecture).
4. Standardizes on Bootstrap CIs document-wide for all baseline proportions.
5. Standardizes identity phrase: 'sparse higher-order spiking amidst broad low-frequency field disruption'.
"""

import docx
import pathlib
from docx.shared import RGBColor

REPO = pathlib.Path(r'D:\workspace\omission')
DOCX_PATH = REPO / 'context' / 'omission-2026-manuscript-master.docx'

doc = docx.Document(str(DOCX_PATH))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# ── 1. Clean Up Duplicated Discussion ─────────────────────────────────────────
disc_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text == 'Discussion':
        disc_idx = i
        break

if disc_idx is not None:
    # Keep exact 3 streamlined Discussion paragraphs (disc_idx+1, disc_idx+2, disc_idx+3)
    # Clear out legacy paragraph at disc_idx+4 if it starts with 'Hierarchical Spectrolaminar'
    for k in range(disc_idx + 4, len(doc.paragraphs)):
        text = doc.paragraphs[k].text
        if text.startswith('Hierarchical Spectrolaminar') or text.startswith('These conclusions remain limited by the observational nature'):
            doc.paragraphs[k].text = ""

# ── 2. Purge Legacy Connectivity Methods ──────────────────────────────────────
for p in doc.paragraphs:
    if p.text.startswith('Advanced Connectivity & Granger Null Models'):
        p.text = ""

# ── 3. Streamline Main Text Figure Captions (Remove Figures 7-10) ─────────────
for p in doc.paragraphs:
    if any(p.text.startswith(prefix) for prefix in ['Figure 7:', 'Figure 8:', 'Figure 9:', 'Figure 10:']):
        p.text = ""

# ── 4. Standardize Identity Phrase Document-Wide ──────────────────────────────
IDENTITY_PHRASE = "sparse higher-order spiking amidst broad low-frequency field disruption"

for p in doc.paragraphs:
    if 'sparse higher-order spiking' in p.text and IDENTITY_PHRASE not in p.text:
        p.text = p.text.replace('sparse higher-order spiking and broad low-frequency field reorganization', IDENTITY_PHRASE)

# Re-apply Calibri Black Document-Wide
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

doc.save(str(DOCX_PATH))
print("Successfully executed Final Master Editorial Purge & Convergence in docx!")

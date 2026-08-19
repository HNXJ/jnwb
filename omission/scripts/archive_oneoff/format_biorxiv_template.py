"""
Formats omission-2026-draft-biorxiv-ready.docx to strictly match template:
D:\workspace\omission\context\omission-2026-draft-format-template.docx

Enforces:
- Font: Calibri for all paragraphs and runs
- Color: 100% solid Black (RGB: 0, 0, 0) for all text and headings
- Saves to both omission-2026-draft-calibrated-formatted.docx and tries omission-2026-draft-biorxiv-ready.docx
"""

import docx
from docx.shared import Pt, RGBColor
import pathlib

REPO = pathlib.Path(r'D:\workspace\omission')
INPUT_DOCX = REPO / 'context' / 'omission-2026-draft-calibrated.docx'
PRIMARY_OUTPUT = REPO / 'context' / 'omission-2026-draft-calibrated-formatted.docx'
TARGET_OUTPUT = REPO / 'context' / 'omission-2026-draft-biorxiv-ready.docx'

doc = docx.Document(str(INPUT_DOCX))

BLACK = RGBColor(0, 0, 0)
FONT_NAME = 'Calibri'

# 1. Update document styles
for s in doc.styles:
    if hasattr(s, 'font') and s.font is not None:
        s.font.name = FONT_NAME
        s.font.color.rgb = BLACK

# 2. Iterate over all paragraphs and runs to strictly set Calibri & Black text
for p in doc.paragraphs:
    for r in p.runs:
        r.font.name = FONT_NAME
        r.font.color.rgb = BLACK

# 3. Iterate over all tables and cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_NAME
                    r.font.color.rgb = BLACK

doc.save(str(PRIMARY_OUTPUT))
print(f"Successfully saved Calibri black formatted manuscript to: {PRIMARY_OUTPUT}")

try:
    doc.save(str(TARGET_OUTPUT))
    print(f"Successfully updated target file: {TARGET_OUTPUT}")
except Exception as e:
    print(f"Note: Could not overwrite {TARGET_OUTPUT} directly due to file lock: {e}")
